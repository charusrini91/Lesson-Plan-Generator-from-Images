from transformers import pipeline
from docx import Document
import datetime
import os

# Initialize models
vision_model = pipeline("image-to-text", model="Salesforce/blip-image-captioning-base")
llm = pipeline("text-generation", model="mistralai/Mistral-7B-Instruct-v0.2", device_map="auto")


def generate_lesson_plan(image_path):
    """
    Step 1: Describe the image
    Step 2: Generate lesson plan text using LLM
    """
    # Generate caption from image
    caption = vision_model(image_path)[0]['generated_text']

    # Prompt for lesson plan
    prompt = f"""
    You are a skilled STEM educator.
    Based on the classroom image described as: "{caption}",
    create a structured lesson plan including the following:
    1. Lesson Title
    2. Objectives
    3. Materials Required
    4. Activities (step-by-step)
    5. Assessment Methods
    Format it clearly with section headers.
    """

    plan_text = llm(prompt, max_new_tokens=400, temperature=0.7)[0]['generated_text']
    return caption, plan_text


def save_to_docx(title, plan_text, output_dir="outputs"):
    """
    Save generated lesson plan to a Word (DOCX) file.
    """
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, f"LessonPlan_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")

    for line in plan_text.split("\n"):
        if line.strip():
            if any(keyword in line.lower() for keyword in ["objective", "material", "activity", "assessment", "lesson"]):
                doc.add_heading(line.strip(), level=1)
            else:
                doc.add_paragraph(line.strip())

    doc.save(file_path)
    return file_path
